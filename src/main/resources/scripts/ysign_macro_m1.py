"""Linux-compatible port of the ySign m1 Word macro.

Replaces placeholders like #s1# in a .docx with named legacy text form
fields (FORMTEXT), padded with spaces — no MS Word required.

Requires: pip install python-docx
Usage:    python ysign_macro_m1.py input.docx [output.docx]
          (defaults to overwriting input in place)

Note: paragraphs containing a placeholder are rebuilt as plain text
(character formatting of the first run is kept; tabs/line breaks are
preserved, but images/other objects in that same paragraph are not).
"""

import copy
import re
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# (placeholder -> (field name, default width in spaces))
FIELDS = {
    **{f"#s{i}#": (f"SIGNA{i:02d}", 30) for i in range(1, 11)},
    **{f"#sd{i}#": (f"DATESIGNEDA{i:02d}", 20) for i in range(1, 11)},
    **{f"#g{i}#": (f"SIGNG{i:02d}", 30) for i in range(1, 8)},
    **{f"#gd{i}#": (f"DATESIGNEDG{i:02d}", 20) for i in range(1, 8)},
    **{f"#sI{i}#": (f"INITIALA{i:02d}", 10) for i in range(1, 11)},
    "#mI1#": ("INITIALM01", 10),
    **{f"#m{i}#": (f"SIGNM{i:02d}", 30) for i in range(1, 5)},
    **{f"#md{i}#": (f"DATESIGNEDM{i:02d}", 20) for i in range(1, 5)},
    **{f"#mn{i}#": (f"NAMEM{i:02d}", 30) for i in range(1, 5)},
}

# Longest placeholders first so #s10# matches before #s1#
PLACEHOLDER_RE = re.compile(
    "|".join(re.escape(p) for p in sorted(FIELDS, key=len, reverse=True))
)

_bookmark_id = [1000]  # unique bookmark ids across the document


def _el(tag, attrs=None):
    e = OxmlElement(tag)
    for k, v in (attrs or {}).items():
        e.set(qn(k), v)
    return e


def _run(rpr, *children):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    for c in children:
        r.append(c)
    return r


def _text_runs(text, rpr):
    """Plain text run(s); '\t' and '\n' become w:tab / w:br."""
    runs = []
    for part in re.split(r"([\t\n])", text):
        if part == "":
            continue
        if part == "\t":
            runs.append(_run(rpr, _el("w:tab")))
        elif part == "\n":
            runs.append(_run(rpr, _el("w:br")))
        else:
            t = _el("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = part
            runs.append(_run(rpr, t))
    return runs


def _field_elements(name, width, rpr):
    """Elements for one named legacy FORMTEXT field (incl. bookmark)."""
    spaces = " " * width
    _bookmark_id[0] += 1
    bid = str(_bookmark_id[0])

    bm_start = _el("w:bookmarkStart", {"w:id": bid, "w:name": name})
    bm_end = _el("w:bookmarkEnd", {"w:id": bid})

    ffdata = _el("w:ffData")
    ffdata.append(_el("w:name", {"w:val": name}))
    ffdata.append(_el("w:enabled"))
    ffdata.append(_el("w:calcOnExit", {"w:val": "0"}))
    text_input = _el("w:textInput")
    text_input.append(_el("w:default", {"w:val": spaces}))
    ffdata.append(text_input)

    begin = _el("w:fldChar", {"w:fldCharType": "begin"})
    begin.append(ffdata)

    instr = _el("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " FORMTEXT "

    sep = _el("w:fldChar", {"w:fldCharType": "separate"})
    end = _el("w:fldChar", {"w:fldCharType": "end"})

    result_t = _el("w:t")
    result_t.set(qn("xml:space"), "preserve")
    result_t.text = spaces

    return [
        bm_start,
        _run(rpr, begin),
        _run(rpr, instr),
        _run(rpr, sep),
        _run(rpr, result_t),
        _run(rpr, end),
        bm_end,
    ]


def _paragraph_text(p_el):
    """Full run text of a paragraph, with tabs/breaks as \t / \n."""
    parts = []
    for r in p_el.findall(qn("w:r")):
        for child in r:
            tag = child.tag
            if tag == qn("w:t"):
                parts.append(child.text or "")
            elif tag == qn("w:tab"):
                parts.append("\t")
            elif tag in (qn("w:br"), qn("w:cr")):
                parts.append("\n")
    return "".join(parts)


def process_paragraph(paragraph):
    """Replace placeholders in one paragraph. Returns fields added."""
    p = paragraph._p
    text = _paragraph_text(p)
    if not PLACEHOLDER_RE.search(text):
        return 0

    runs = p.findall(qn("w:r"))
    rpr = runs[0].find(qn("w:rPr")) if runs else None
    for r in runs:
        p.remove(r)

    new_elements, pos, count = [], 0, 0
    for m in PLACEHOLDER_RE.finditer(text):
        if m.start() > pos:
            new_elements.extend(_text_runs(text[pos : m.start()], rpr))
        name, width = FIELDS[m.group()]
        new_elements.extend(_field_elements(name, width, rpr))
        count += 1
        pos = m.end()
    if pos < len(text):
        new_elements.extend(_text_runs(text[pos:], rpr))

    for e in new_elements:
        p.append(e)
    return count


def _iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def m1(in_path, out_path=None):
    doc = Document(in_path)
    total = 0
    for paragraph in _iter_paragraphs(doc):
        total += process_paragraph(paragraph)
    for section in doc.sections:
        for part in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            for paragraph in _iter_paragraphs(part):
                total += process_paragraph(paragraph)
    doc.save(out_path or in_path)
    print(f"Done: {total} form fields added -> {out_path or in_path}")


if __name__ == "__main__":
    if len(sys.argv) not in (2, 3):
        sys.exit("Usage: python ysign_macro_m1.py input.docx [output.docx]")
    m1(sys.argv[1], sys.argv[2] if len(sys.argv) == 3 else None)
